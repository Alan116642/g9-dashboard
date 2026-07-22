"""Generate source-backed charts, notebooks, Word/PDF-ready reports, and W5 copy."""
from __future__ import annotations

import json
import math
from pathlib import Path
from typing import Any

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import nbformat as nbf
import numpy as np
import pandas as pd
from matplotlib.patches import Patch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from g9_pipeline import CHART_DIR, PRIVATE_DIR, REPORTS_DIR, ROOT, CleanBundle, json_ready
from visual_catalog import (
    WHITEPAPER_CHAPTER_STRUCTURE,
    WHITEPAPER_CHART_SEQUENCE,
    WHITEPAPER_TASK_CHART_KEYS,
    WHITEPAPER_TASK_SECTIONS,
    build_extended_charts,
)
from visual_semantics import (
    GOOD,
    GOOD_GREEN,
    IMPROVE,
    IMPROVE_RED,
    MUTED_GRAY,
    NEUTRAL_BLUE,
    STATUS_COLORS,
    WATCH,
    WATCH_AMBER,
    colors_for,
)

plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "Arial Unicode MS", "DejaVu Sans"]
plt.rcParams["axes.unicode_minus"] = False

BLUE = NEUTRAL_BLUE.lstrip("#")
DARK_BLUE = "1F4D78"
INK = "172033"
MUTED = "64748B"
LIGHT = "F2F4F7"
GOLD = WATCH_AMBER.lstrip("#")


def _save(fig: plt.Figure, name: str) -> Path:
    CHART_DIR.mkdir(parents=True, exist_ok=True)
    path = CHART_DIR / name
    fig.savefig(path, dpi=220, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    return path


def build_charts(charts: dict[str, pd.DataFrame], bundle: CleanBundle, analysis: dict[str, Any]) -> dict[str, Path]:
    output: dict[str, Path] = {}
    colors = {"首次面谈": "#D97706", "其他方式": "#2563EB"}

    fig, ax = plt.subplots(figsize=(10, 5.5))
    for label, group in charts["propensity"].groupby("处理组"):
        ax.hist(group["倾向得分"], bins=35, alpha=0.55, density=True, color=colors[label], label=f"{label}（样本量={len(group):,}）")
    ax.set(title="倾向得分共同支持", xlabel="倾向得分", ylabel="密度")
    ax.legend(frameon=False); ax.spines[["top", "right"]].set_visible(False)
    output["propensity"] = _save(fig, "w2_01_propensity_overlap.png")

    balance = charts["balance"].copy()
    balance["排序"] = balance[["匹配前SMD", "匹配后SMD"]].abs().max(axis=1)
    balance = balance.sort_values("排序").tail(18)
    fig, ax = plt.subplots(figsize=(10, 7))
    y = np.arange(len(balance))
    ax.scatter(balance["匹配前SMD"].abs(), y, label="匹配前", color="#94A3B8", s=40)
    ax.scatter(balance["匹配后SMD"].abs(), y, label="匹配后（通过）", color=GOOD_GREEN, s=45)
    ax.axvline(0.1, color=IMPROVE_RED, linestyle="--", linewidth=1.3, label="需改进阈值：绝对标准化均值差=0.1")
    ax.set_yticks(y, balance["协变量"]); ax.set(xlabel="绝对标准化均值差", title="匹配前后协变量平衡")
    ax.legend(frameon=False); ax.spines[["top", "right"]].set_visible(False)
    output["balance"] = _save(fig, "w2_02_love_plot.png")

    effects = charts["effects"].copy()
    effects["方法"] = effects["方法"].replace({
        "PSM ATT": "倾向得分匹配的处理组平均效应",
        "AIPW ATT": "增强逆概率加权的处理组平均效应",
    })
    fig, ax = plt.subplots(figsize=(9, 4.6))
    y = np.arange(len(effects))
    ax.errorbar(effects["估计"], y, xerr=[effects["估计"] - effects["下限"], effects["上限"] - effects["估计"]], fmt="o", color=WATCH_AMBER, ecolor=MUTED_GRAY, capsize=5, label="需要关注：区间跨零")
    ax.axvline(0, color="#111827", linestyle="--", linewidth=1)
    ax.set_yticks(y, effects["方法"]); ax.set(xlabel="下订概率差", title="跟进方式效应估计与 95% 置信区间")
    ax.legend(frameon=False); ax.spines[["top", "right"]].set_visible(False)
    output["effects"] = _save(fig, "w2_03_followup_effects.png")

    sensitivity = charts["sensitivity"]
    fig, ax = plt.subplots(figsize=(9, 4.8))
    ax.plot(sensitivity["Gamma"], sensitivity["p值上界"], color="#2563EB", linewidth=2.2)
    ax.axhline(0.05, color="#B91C1C", linestyle="--", linewidth=1.2)
    ax.set(xlabel="隐藏偏差赔率上界", ylabel="双侧检验概率上界", title="匹配对隐藏偏差敏感性")
    ax.set_ylim(0, 1.02); ax.spines[["top", "right"]].set_visible(False)
    output["sensitivity"] = _save(fig, "w2_04_rosenbaum_bounds.png")

    delivery = charts["delivery"]
    score_rows = delivery.loc[delivery["结果"].str.startswith("评分")].copy()
    complaint_rows = delivery.loc[delivery["结果"].str.startswith("投诉")].copy()
    fig, axes = plt.subplots(1, 2, figsize=(11, 4.8))
    for ax, frame, title, xlabel, scale in [
        (axes[0], score_rows, "交付评分差异", "评分差异（分）", 1.0),
        (axes[1], complaint_rows, "投诉概率调整后关联", "概率差异（百分点）", 100.0),
    ]:
        y = np.arange(len(frame))
        estimate = (frame["估计"] * scale).to_numpy(dtype=float)
        lower = (frame["下限"] * scale).to_numpy(dtype=float)
        upper = (frame["上限"] * scale).to_numpy(dtype=float)
        ax.errorbar(estimate, y, xerr=[estimate - lower, upper - estimate], fmt="o", color=WATCH_AMBER, ecolor=MUTED_GRAY, capsize=5)
        ax.axvline(0, color="#111827", linestyle="--", linewidth=1)
        ax.set_yticks(y, frame["结果"]); ax.set(xlabel=xlabel, title=title)
        ax.spines[["top", "right"]].set_visible(False)
    fig.legend(handles=[Patch(color=WATCH_AMBER, label="需要关注：置信区间跨零")], frameon=False, loc="upper center", bbox_to_anchor=(0.5, 0.94))
    fig.suptitle("延迟交付与结果变量的估计（不同单位分面展示）")
    fig.tight_layout()
    output["delivery"] = _save(fig, "w2_05_delivery_associations.png")

    channel = charts["channel_descriptive"].sort_values("转化率")
    fig, ax = plt.subplots(figsize=(9, 5))
    channel_colors = colors_for(channel["转化率"], "higher")
    for idx, (_, row) in enumerate(channel.iterrows()):
        ax.errorbar(
            row["转化率"] * 100,
            idx,
            xerr=[[max(0.0, (row["转化率"] - row["置信区间下限"]) * 100)], [max(0.0, (row["置信区间上限"] - row["转化率"]) * 100)]],
            fmt="o",
            color=channel_colors[idx],
            ecolor=MUTED_GRAY,
            capsize=4,
        )
    ax.set_yticks(np.arange(len(channel)), channel["渠道"]); ax.set(xlabel="转化率（%）", title="渠道转化率及威尔逊 95% 置信区间")
    ax.legend(handles=[Patch(color=STATUS_COLORS[label], label=label) for label in (IMPROVE, WATCH, GOOD)], frameon=False)
    ax.spines[["top", "right"]].set_visible(False)
    output["channel"] = _save(fig, "w2_06_channel_conversion.png")

    standardized = charts["channel_standardized"].sort_values("标准化边际转化率")
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.barh(standardized["渠道"], standardized["标准化边际转化率"] * 100, color=colors_for(standardized["标准化边际转化率"], "higher"))
    ax.set(xlabel="标准化边际转化率（%）", title="渠道调整后边际转化率（关联性）")
    ax.legend(
        handles=[Patch(color=STATUS_COLORS[label], label=label) for label in (IMPROVE, WATCH, GOOD)],
        frameon=False,
        loc="center left",
        bbox_to_anchor=(1.01, 0.5),
    )
    ax.spines[["top", "right"]].set_visible(False)
    output["channel_standardized"] = _save(fig, "w2_07_channel_standardized.png")

    importance = charts["model_importance"].sort_values("预测重要性")
    fig, ax = plt.subplots(figsize=(9, 5))
    ax.barh(importance["特征"], importance["预测重要性"], color=["#2563EB" if v >= 0 else "#CBD5E1" for v in importance["预测重要性"]])
    ax.axvline(0, color="#111827", linewidth=0.8)
    ax.set(xlabel="时间外受试者工作特征曲线下面积的置换变化", title="基线预测特征重要性（非因果）")
    ax.spines[["top", "right"]].set_visible(False)
    output["importance"] = _save(fig, "w3_01_predictive_importance.png")

    strategy = charts["strategy"].head(12).sort_values("优先级分")
    labels = strategy["城市"] + " / " + strategy["渠道"]
    fig, ax = plt.subplots(figsize=(10, 6))
    ax.barh(labels, strategy["优先级分"], color=colors_for(strategy["优先级分"], "lower"))
    ax.set(xlabel="机会优先级分（不是预算投资回报率）", title="城市 × 渠道运营机会排序")
    ax.legend(handles=[Patch(color=STATUS_COLORS[label], label=label) for label in (IMPROVE, WATCH, GOOD)], frameon=False)
    ax.spines[["top", "right"]].set_visible(False)
    output["strategy"] = _save(fig, "w4_01_opportunity_priority.png")

    output.update(build_extended_charts(bundle, analysis, CHART_DIR, strategy=charts["strategy"]))

    chart_map = {
        key: {
            "path": str(path.relative_to(ROOT)),
            "source": "data/processed/analysis_results.json and derived private aggregates",
            "purpose": key,
            "palette_policy": (
                "red=优先改进; amber=需要关注; green=表现较好; "
                "blue/gray=描述性或不可判优劣; 红橙绿为图内相对分位，非目标、显著性或因果效应"
            ),
        }
        for key, path in output.items()
    }
    for position, spec in enumerate(WHITEPAPER_CHART_SEQUENCE, start=1):
        chart_map[spec["key"]].update({
            "whitepaper_order": position,
            "section": spec["section"],
            "caption": spec["caption"],
            "interpretation": spec["note"],
        })
    (REPORTS_DIR / "chart_map.json").write_text(json.dumps(chart_map, ensure_ascii=False, indent=2), encoding="utf-8")
    return output


def _set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}"); tcMar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def _set_table_geometry(table, widths: list[int]) -> None:
    table.autofit = False; table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tblPr = table._tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW"); tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths))); tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd"); tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120"); tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid): grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tcW = cell._tc.get_or_add_tcPr().get_or_add_tcW(); tcW.set(qn("w:w"), str(width)); tcW.set(qn("w:type"), "dxa")
            _set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_run(run, size=11, bold=False, color=INK, font="Calibri", italic=False):
    run.font.name = font; run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font); run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size); run.bold = bold; run.italic = italic; run.font.color.rgb = RGBColor.from_string(color)


def _page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 "); _set_run(run, size=9, color=MUTED)
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), "PAGE"); paragraph._p.append(fld)
    run2 = paragraph.add_run(" 页"); _set_run(run2, size=9, color=MUTED)


def _setup_doc(title: str, subtitle: str, report_type: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5); section.page_height = Inches(11)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    section.header_distance = section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"; normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11); normal.paragraph_format.space_after = Pt(6); normal.paragraph_format.line_spacing = 1.10
    for style_name, size, color, before, after in (("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",12,DARK_BLUE,8,4)):
        style = doc.styles[style_name]; style.font.name = "Calibri"; style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri"); style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size); style.font.bold = True; style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before); style.paragraph_format.space_after = Pt(after)
    header = section.header.paragraphs[0]; header.clear(); header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_run(header.add_run(report_type), size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]; footer.clear(); _page_number(footer)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
    _set_run(p.add_run(report_type.upper()), size=11, bold=True, color=GOLD)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    _set_run(p.add_run(title), size=23, bold=True, color="000000")
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(14)
    _set_run(p.add_run(subtitle), size=14, color="374151")
    for label, value in (("对象", "G9 销售运营决策团队"), ("数据截止", "2025-12-31（交付数据延伸至 2026-02-27）"), ("证据原则", "不以显著性为目标；诊断失败即降级结论")):
        p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(2)
        _set_run(p.add_run(f"{label}："), bold=True, color="000000"); _set_run(p.add_run(value), color="000000")
    doc.add_paragraph()
    return doc


def _add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(6); p.paragraph_format.left_indent = Inches(0.5); p.paragraph_format.first_line_indent = Inches(-0.25)
        _set_run(p.add_run(item), color="111827")


def _add_image(doc: Document, path: Path, caption: str):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(6.15))
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER; cap.paragraph_format.space_after = Pt(8)
    _set_run(cap.add_run(caption), size=9, color=MUTED, italic=True)


def _add_result_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["分析", "估计", "95% 置信区间", "证据等级"]
    for i, value in enumerate(headers):
        cell = table.rows[0].cells[i]; cell.text = value
        cell._tc.get_or_add_tcPr().append(OxmlElement("w:shd")); cell._tc.tcPr[-1].set(qn("w:fill"), LIGHT)
        for run in cell.paragraphs[0].runs: _set_run(run, size=10, bold=True, color="111827")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for run in cells[i].paragraphs[0].runs: _set_run(run, size=9.5, color="111827")
    _set_table_geometry(table, [2200, 1500, 2700, 2960])
    doc.add_paragraph()


def _add_color_legend(doc: Document) -> None:
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(8)
    _set_run(p.add_run("■ 优先改进  "), size=10, bold=True, color=IMPROVE_RED.lstrip("#"))
    _set_run(p.add_run("■ 需要关注  "), size=10, bold=True, color=WATCH_AMBER.lstrip("#"))
    _set_run(p.add_run("■ 表现较好  "), size=10, bold=True, color=GOOD_GREEN.lstrip("#"))
    _set_run(p.add_run("■ 描述性/不可判优劣"), size=10, bold=True, color=NEUTRAL_BLUE.lstrip("#"))
    p = doc.add_paragraph()
    _set_run(p.add_run("红橙绿按同一张图中的相对分位识别改进优先级，不是业务目标、显著性或因果效应；蓝灰图只用于描述结构。"), size=9.5, color=MUTED)


def _pp(value: float | None) -> str:
    return "不可估计" if value is None else f"{value * 100:+.2f} 个百分点"


def _ci_pp(ci: list[float | None]) -> str:
    if ci[0] is None: return "不可估计"
    return f"{ci[0] * 100:+.2f} 至 {ci[1] * 100:+.2f} 个百分点"


def generate_w2_docx(analysis: dict[str, Any], audit: dict[str, Any], images: dict[str, Path]) -> Path:
    psm = analysis["psm"]["psm_att"]; aipw = analysis["psm"]["aipw_att"]
    score = analysis["delivery"]["delivery_score_adjusted_association"]; complaint = analysis["delivery"]["complaint_adjusted_association"]
    doc = _setup_doc("G9 因果归因与证据审计", "跟进方式、交付延迟与渠道差异的重新分析", "W2 分析报告")
    doc.add_heading("执行摘要", level=1)
    _add_bullets(doc, [
        f"匹配后协变量平衡通过，但首次面谈的 ATT 为 {_pp(psm['estimate'])}，95% CI 为 {_ci_pp(psm['ci95'])}；未发现可靠的转化增益证据。",
        f"AIPW 稳健性结果为 {_pp(aipw['estimate'])}，置信区间同样跨零，方向与 PSM 不构成可行动的正向证据。",
        f"延迟交付与交付评分的调整后差异为 {score['estimate']:+.3f} 分，95% CI {score['ci95'][0]:+.3f} 至 {score['ci95'][1]:+.3f}；不能支持评分存在显著下降。",
        "当前数据没有合法 RDD 断点，也没有多触点渠道路径或渠道级花费；RDD、Shapley 多触点归因及渠道 ROI 均不具备识别条件。",
    ])
    doc.add_heading("交付数据需要先修复粒度", level=1)
    p = doc.add_paragraph(); _set_run(p.add_run(f"5,000 条交付记录仅对应 {audit['row_counts']['unique_orders']:,} 个唯一订单；{audit['critical_conflict_orders']:,} 个订单存在关键字段冲突。主要交付分析仅使用 {audit['usable_orders']:,} 个无关键冲突订单。"))
    doc.add_heading("面谈跟进没有可靠的正向增益证据", level=1)
    _add_result_table(doc, [["PSM ATT", _pp(psm["estimate"]), _ci_pp(psm["ci95"]), psm["evidence_level"]], ["AIPW ATT", _pp(aipw["estimate"]), _ci_pp(aipw["ci95"]), aipw["evidence_level"]]])
    p = doc.add_paragraph(); _set_run(p.add_run(f"共匹配 {analysis['psm']['diagnostics']['matched_pairs']:,} 对线索；匹配后最大绝对 SMD 为 {analysis['psm']['diagnostics']['max_abs_smd_after']:.3f}。平衡达标只说明观察协变量更可比，不会把跨零结果变成显著效果。"))
    _add_image(doc, images["propensity"], "图 1. 倾向得分共同支持")
    _add_image(doc, images["balance"], "图 2. 匹配前后协变量平衡")
    _add_image(doc, images["effects"], "图 3. PSM 与 AIPW 效应估计")
    _add_image(doc, images["sensitivity"], "图 4. 匹配对隐藏偏差敏感性")
    doc.add_heading("延迟交付只能做调整后关联分析", level=1)
    p = doc.add_paragraph(); _set_run(p.add_run("评分模型控制城市、配置和交付月份，并采用 HC3 稳健标准误。评分差异置信区间跨零，因此不能得出负向影响；投诉结果虽出现统计关联，但经 FDR 后仍只能作为非因果线索，需补充承诺交付日、实际交付日和延迟原因。"))
    _add_result_table(doc, [["评分调整后关联", f"{score['estimate']:+.3f} 分", f"{score['ci95'][0]:+.3f} 至 {score['ci95'][1]:+.3f}", score["evidence_level"]], ["投诉概率调整后关联", _pp(complaint["estimate"]), _ci_pp(complaint["ci95"]), complaint["evidence_level"]]])
    _add_image(doc, images["delivery"], "图 5. 延迟交付与评分/投诉的调整后关联")
    doc.add_heading("RDD 可行性审计：不满足识别条件", level=1)
    _add_bullets(doc, ["没有连续分配变量及事先确定的阈值。", "交付里程中位数不是延迟交付的业务规则，不能被用作断点。", "因此不再生成 McCrary、带宽稳健性或安慰剂 RDD 结论。"])
    doc.add_heading("渠道输出改为描述和标准化关联", level=1)
    p = doc.add_paragraph(); _set_run(p.add_run("每条线索只有一个获客渠道，无法计算多触点 Shapley；市场费用只到城市-月份粒度，也无法计算渠道 ROI。报告保留渠道转化率及标准化边际转化率，并明确标注为关联性结果。"))
    _add_image(doc, images["channel"], "图 6. 渠道转化率及 95% 置信区间")
    _add_image(doc, images["channel_standardized"], "图 7. 渠道标准化边际转化率")
    doc.add_heading("Recommended Next Steps", level=1)
    _add_bullets(doc, ["补采下订时间戳，保证处理先于结果。", "补采承诺交付日、实际交付日及延迟原因，重新设计交付效果识别。", "补采渠道级花费和多触点路径后，再启用 ROI 或 Shapley。", "在面谈策略上优先进行随机或准随机试验，不以当前跨零估计作为扩张依据。"])
    doc.add_heading("Further Questions", level=1)
    _add_bullets(doc, ["首次跟进方式是否由销售员自主选择，还是由门店规则分配？", "下订发生在首次跟进之前的比例是多少？", "冲突交付记录来自更新历史、系统合并还是数据生成逻辑？"])
    doc.add_heading("Caveats and Assumptions", level=1)
    _add_bullets(doc, psm["assumptions"] + score["assumptions"])
    path = REPORTS_DIR / "W2_causal_attribution_report.docx"; doc.save(path); return path


def _whitepaper_section_copy(title: str, analysis: dict[str, Any], audit: dict[str, Any]) -> list[str]:
    psm = analysis["psm"]["psm_att"]
    model = analysis["model"]["metrics"]
    copy = {
        "1.1 问题与答案": [
            "业务目标是在不增加总预算的前提下推动销量改善。当前数据可以识别数据质量、流程和运营机会，但不能直接证明能够实现 15% 增长；增长目标必须由后续试验和完整成本数据验证。",
        ],
        "1.2 五大核心发现": [
            f"交付表中 {audit['critical_conflict_orders']:,}/{audit['row_counts']['unique_orders']:,} 个唯一订单存在关键字段冲突；首次面谈效应为 {_pp(psm['estimate'])} 且置信区间跨零；时间外模型的受试者工作特征曲线下面积为 {model['roc_auc']:.3f}；渠道缺少可分配花费；公开看板只发布脱敏汇总。",
        ],
        "1.3 目标与证据边界": [
            "本白皮书区分描述性结果、调整后关联、条件因果估计和预测解释。没有合法断点、渠道级成本、多触点路径或有效生存时间时，对应方法停止使用并在原章节位置保留可行性审计。",
        ],
        "2.1 转化漏斗": ["统一使用线索级和无关键冲突订单级口径，避免多表连接造成线索、订单或投诉重复计数。"],
        "2.2 渠道效率全景": ["渠道图展示规模、每百线索订单数和跟进覆盖；它们是运营比较，不是投资回报率或因果贡献。"],
        "2.3 城市效率地图": ["城市差异用于发现需要诊断的区域，不直接推断门店或销售人员能力。"],
        "2.4 时间趋势": ["所有日期统一按月展示；交付数据延伸至 2026-02-27，跨期比较需注意不同业务时间窗。"],
        "2.5 客户画像与试驾行为": ["客户画像和试驾行为仅用于描述与预测，不作为差异化待遇或因果结论。"],
        "2.6 数据质量与时序审计": ["关键冲突订单退出主要交付分析；入职前跟进、交付后跟进和同日多种首次方式均被显式审计。"],
        "3.1 为什么需要因果推断": ["跟进方式由业务选择而非随机分配，原始转化差异可能来自城市、月份、渠道或客户构成。"],
        "3.2 倾向得分匹配：跟进方式的条件因果估计": [f"匹配后协变量平衡通过，但处理组平均效应为 {_pp(psm['estimate'])}，95% 置信区间 {_ci_pp(psm['ci95'])}，未发现可靠增益证据；且缺少下订时间戳，时序假设无法完全验证。"],
        "3.3 交付延迟：调整后关联与断点可行性审计": ["数据没有连续分配变量和预先规定阈值，因此断点回归不可识别。延迟交付只报告对评分与投诉的调整后关联。"],
        "3.4 渠道归因：描述、标准化与识别边界": ["每条线索只有一个获客渠道，无法计算多触点沙普利归因；市场费用只有城市-月份总额，无法计算渠道投资回报率或单均成本。"],
        "3.5 异质性与分层线索": ["城市-渠道分层仅用于提出进一步诊断和试验假设，多重比较不应被解读为稳定的异质性因果效应。"],
        "4.1 转化预测基线": [f"模型采用时间顺序训练与验证，只使用线索创建时可得字段；时间外受试者工作特征曲线下面积为 {model['roc_auc']:.3f}，只保留为验证基线。"],
        "4.2 关键预测因子解读": ["特征重要性与行为分布用于解释预测和设计补采方案，不代表变量对转化的因果贡献。"],
        "4.3 聚合风险预警": ["由于没有合法事件时间定义，生存分析已停用；风险预警改为城市、配置和月份层面的聚合监控，线上不公开客户级预测。"],
        "5.1 机会优先级与资源配置": ["缺少渠道级花费时，优化降级为机会排序；优先级综合规模、转化、覆盖和风险，只用于确定诊断与试验顺序。"],
        "5.2 定价博弈可行性审计": ["当前数据没有竞品价格、价格弹性、成交价或对手反应，因此不生成虚假的博弈均衡或最优定价。后续补齐数据后再建情景模型。"],
        "5.3 人效与售后改善线索": ["当前缺少班次、在岗时段、线索容量和完整分配约束，不能求解最优排班；现阶段只监控跟进覆盖和售后风险。"],
        "5.4 综合策略推荐": ["先修复数据基础，再运行跟进方式随机化试点，并建立渠道成本与多触点采集；只有增量效果和成本同时可信时才进入预算优化。"],
        "6.1 0-30 天：口径与数据责任": ["锁定线索、订单、跟进事件唯一键，建立冲突订单修复责任、时间戳口径和公开数据抑制检查。"],
        "6.2 31-60 天：试验与补采": ["开展跟进方式随机化试点，补采下订、承诺交付、实际交付、渠道花费、触点路径、竞品价格和排班约束。"],
        "6.3 61-90 天：规则更新与经营复盘": ["根据试验结果更新策略规则，把数据哈希、样本量、置信区间和证据等级接入月度经营复盘。"],
    }
    return copy[title]


def _add_strategy_table(doc: Document, strategy: pd.DataFrame) -> None:
    table = doc.add_table(rows=1, cols=4); table.style = "Table Grid"
    for i, value in enumerate(["城市 / 渠道", "线索数", "转化率", "建议"]):
        table.rows[0].cells[i].text = value
        for run in table.rows[0].cells[i].paragraphs[0].runs: _set_run(run, size=10, bold=True)
    for _, row in strategy.head(8).iterrows():
        values = [f"{row['城市']} / {row['渠道']}", f"{row['线索数']:,}", f"{row['转化率']:.1%}", row["建议"]]
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = value
            for run in cells[i].paragraphs[0].runs: _set_run(run, size=9)
    _set_table_geometry(table, [2300, 1300, 1500, 4260])


def generate_whitepaper(analysis: dict[str, Any], audit: dict[str, Any], images: dict[str, Path], strategy: pd.DataFrame) -> tuple[Path, Path]:
    doc = _setup_doc("G9 销售运营优化白皮书", "按参考文档章节重构的证据审计版 · 54 张全中文图表", "最终白皮书")
    p = doc.add_paragraph(); _set_run(p.add_run("结构说明："), size=10.5, bold=True); _set_run(p.add_run("沿用原白皮书的执行摘要、数据诊断、因果归因、预测预警、策略优化、实施路线图和方法附录结构；所有结论替换为当前可复现结果。"), size=10.5)
    _add_color_legend(doc)
    specs = {item["key"]: item for item in WHITEPAPER_CHART_SEQUENCE}
    figure_no = 1
    for chapter_index, chapter in enumerate(WHITEPAPER_CHAPTER_STRUCTURE):
        if chapter_index:
            doc.add_page_break()
        doc.add_heading(chapter["chapter"], level=1)
        meta = WHITEPAPER_TASK_SECTIONS[chapter["task"]]
        p = doc.add_paragraph(); _set_run(p.add_run("对应任务："), size=9.5, bold=True, color="2E74B5"); _set_run(p.add_run(f"{chapter['task']} · {meta['state']}"), size=9.5, color="475569")
        for section in chapter["sections"]:
            doc.add_heading(section["title"], level=2)
            for paragraph in _whitepaper_section_copy(section["title"], analysis, audit):
                p = doc.add_paragraph(); _set_run(p.add_run(paragraph), size=10.5)
            for key in section["keys"]:
                spec = specs[key]
                p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(4)
                _set_run(p.add_run(spec["note"]), size=10.5, color="111827")
                _add_image(doc, images[key], f"图 {figure_no}. {spec['caption']}")
                figure_no += 1
            if section["title"] == "5.4 综合策略推荐":
                doc.add_page_break()
                _add_strategy_table(doc, strategy)

    doc.add_page_break()
    doc.add_heading("附录  方法论与技术栈", level=1)
    doc.add_heading("A.1 五周任务交付总览", level=2)
    task_table = doc.add_table(rows=1, cols=4); task_table.style = "Table Grid"
    for i, value in enumerate(["任务", "图表数", "完成状态", "主要输出"]):
        task_table.rows[0].cells[i].text = value
        for run in task_table.rows[0].cells[i].paragraphs[0].runs: _set_run(run, size=9.5, bold=True)
    for task, meta in WHITEPAPER_TASK_SECTIONS.items():
        cells = task_table.add_row().cells
        values = [task.split(" ", 1)[0], str(len(WHITEPAPER_TASK_CHART_KEYS[task])), meta["state"], meta["deliverable"]]
        for i, value in enumerate(values):
            cells[i].text = value
            for run in cells[i].paragraphs[0].runs: _set_run(run, size=8.5)
    _set_table_geometry(task_table, [900, 900, 2300, 5260])
    doc.add_heading("A.2 方法与复现", level=2)
    _add_bullets(doc, ["数据处理：Pandas、NumPy；原始 Excel 保持只读。", "因果分析：基线协变量倾向得分匹配、共同支持、卡钳、城市/月约束、增强逆概率加权与重采样。", "预测：时间顺序训练/验证；仅使用处理前可得变量。", "看板：Streamlit；线上只读取样本量阈值为 10 的脱敏汇总。"])
    doc.add_heading("A.3 限制与后续问题", level=2)
    _add_bullets(doc, ["缺少下订、承诺交付和实际交付完整时间戳。", "缺少渠道级花费、多触点路径、竞品价格、成交价和排班约束。", "预测重要性不是因果贡献；机会优先级不是投资回报率。", "原始工作簿未被修改。"])
    docx_path = REPORTS_DIR / "G9_销售运营优化白皮书_证据审计版.docx"; doc.save(docx_path)

    md = ["# G9 销售运营优化白皮书", "", f"> 按参考文档章节重构的证据审计版，共 {len(WHITEPAPER_CHART_SEQUENCE)} 张全中文图表。", "", "颜色含义：红=优先改进，橙=需要关注，绿=表现较好，蓝灰=描述性/不可判优劣。", ""]
    figure_no = 1
    for chapter in WHITEPAPER_CHAPTER_STRUCTURE:
        md += [f"## {chapter['chapter']}", "", f"**对应任务：** {chapter['task']} · {WHITEPAPER_TASK_SECTIONS[chapter['task']]['state']}", ""]
        for section in chapter["sections"]:
            md += [f"### {section['title']}", ""]
            md += _whitepaper_section_copy(section["title"], analysis, audit) + [""]
            for key in section["keys"]:
                spec = specs[key]
                md.append(f"{figure_no}. **{spec['caption']}**：{spec['note']}")
                figure_no += 1
            md.append("")
    md += ["## 附录  方法论与技术栈", "", "- 原始 Excel 只读；分析结果、图表、看板和报告使用同一结果接口。", "- 不具备识别条件的方法保留可行性审计，不生成伪结论。", "- 详细 W1-W5 交付状态见项目交付物总验收清单。", ""]
    md_path = REPORTS_DIR / "W5_white_paper.md"; md_path.write_text("\n".join(md), encoding="utf-8")
    return docx_path, md_path


def generate_notebook() -> Path:
    notebook = nbf.v4.new_notebook()
    notebook["metadata"]["kernelspec"] = {"display_name": "Python 3", "language": "python", "name": "python3"}
    notebook["cells"] = [
        nbf.v4.new_markdown_cell("# G9 因果归因重新分析\n\n## tl;dr\n\n本 Notebook 从不可修改的原始 Excel 重建清洗层与因果结果。结论只以执行输出为准。"),
        nbf.v4.new_markdown_cell("## Context & Methods\n\n### Key Assumptions\n\n- 下订时间戳缺失，因此跟进处理先于下订是不可验证假设。\n- PSM 只使用基线协变量，并以匹配后 SMD 判断平衡。\n- 交付记录关键字段冲突订单不进入主要估计。\n- RDD、多触点 Shapley 和渠道 ROI 在当前数据下不具备识别条件。"),
        nbf.v4.new_code_cell("from pathlib import Path\nimport json, sys\nROOT = Path.cwd()\nif not (ROOT / 'src').exists():\n    ROOT = ROOT.parent\nsys.path.insert(0, str(ROOT / 'src'))\nfrom g9_pipeline import build_clean_data, run_analysis\nbundle = build_clean_data(save=False)\nanalysis, charts = run_analysis(bundle, publish=False)\nprint('analysis generated:', analysis['generated_at'])"),
        nbf.v4.new_markdown_cell("## Data\n\n确认分析粒度、重复订单和时序异常。"),
        nbf.v4.new_code_cell("import pandas as pd\npd.DataFrame({\n    '表': ['线索级', '订单级', '跟进事件'],\n    '行数': [len(bundle.leads), len(bundle.orders), len(bundle.followups)],\n    '粒度': ['线索ID唯一', '订单ID唯一', '跟进ID事件']\n})"),
        nbf.v4.new_code_cell("pd.Series(bundle.audit['temporal_issues'], name='数量').to_frame()"),
        nbf.v4.new_markdown_cell("## Results\n\n### 跟进方式 PSM 与 AIPW"),
        nbf.v4.new_code_cell("pd.DataFrame([analysis['psm']['psm_att'], analysis['psm']['aipw_att']], index=['PSM ATT','AIPW ATT'])[['estimate','ci95','p_value','sample_size','balance_status','evidence_level']]"),
        nbf.v4.new_code_cell("charts['balance'].sort_values('匹配后SMD', key=lambda s: s.abs(), ascending=False).head(15)"),
        nbf.v4.new_markdown_cell("### 交付延迟与结果"),
        nbf.v4.new_code_cell("pd.DataFrame(analysis['delivery']).T[['estimate','ci95','p_value','sample_size','method','evidence_level']]"),
        nbf.v4.new_markdown_cell("### 渠道描述与标准化关联"),
        nbf.v4.new_code_cell("pd.DataFrame(analysis['channel']['descriptive'])"),
        nbf.v4.new_code_cell("pd.DataFrame(analysis['channel']['standardized'])"),
        nbf.v4.new_markdown_cell("## Takeaways\n\n- 匹配平衡通过，但面谈 ATT 的 95% 置信区间跨零。\n- 延迟交付对评分的调整后差异跨零。\n- 当前数据不能支持 RDD、多触点 Shapley 或渠道 ROI。\n- 所有结果应连同假设与证据等级一起使用。"),
    ]
    path = REPORTS_DIR / "W2_causal_attribution.ipynb"
    nbf.write(notebook, path)
    return path


def _pdf_styles():
    from reportlab.lib.colors import HexColor
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    try:
        pdfmetrics.registerFont(TTFont("MSYH", str(font_path), subfontIndex=0))
        font_name = "MSYH"
    except Exception:
        font_name = "Helvetica"
    base = getSampleStyleSheet()
    styles = {
        "title": ParagraphStyle("G9Title", parent=base["Title"], fontName=font_name, fontSize=24, leading=31, textColor=HexColor("#172033"), spaceAfter=18),
        "h1": ParagraphStyle("G9H1", parent=base["Heading1"], fontName=font_name, fontSize=16, leading=22, textColor=HexColor("#2E74B5"), spaceBefore=14, spaceAfter=8),
        "body": ParagraphStyle("G9Body", parent=base["BodyText"], fontName=font_name, fontSize=10.5, leading=17, textColor=HexColor("#172033"), spaceAfter=8),
        "bullet": ParagraphStyle("G9Bullet", parent=base["BodyText"], fontName=font_name, fontSize=10.5, leading=17, leftIndent=16, firstLineIndent=-8, spaceAfter=6),
        "caption": ParagraphStyle("G9Caption", parent=base["BodyText"], fontName=font_name, fontSize=8.5, leading=12, textColor=HexColor("#64748B"), alignment=1, spaceAfter=10),
    }
    return font_name, styles


def _pdf_page(canvas, doc):
    canvas.saveState(); canvas.setFont("Helvetica", 8); canvas.setFillColorRGB(0.39, 0.45, 0.55)
    canvas.drawRightString(doc.pagesize[0] - 54, 28, f"第 {doc.page} 页"); canvas.restoreState()


def generate_pdf_reports(analysis: dict[str, Any], audit: dict[str, Any], images: dict[str, Path], strategy: pd.DataFrame) -> tuple[Path, Path]:
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_CENTER
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.platypus import Image as RLImage
    from reportlab.platypus import KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    font_name, styles = _pdf_styles()
    psm = analysis["psm"]["psm_att"]; aipw = analysis["psm"]["aipw_att"]
    score = analysis["delivery"]["delivery_score_adjusted_association"]

    def image_block(key: str, caption: str):
        return [RLImage(str(images[key]), width=6.15 * inch, height=3.55 * inch), Paragraph(caption, styles["caption"])]

    def compact_image_block(key: str, caption: str, note: str):
        return KeepTogether([
            Paragraph(note, styles["body"]),
            RLImage(str(images[key]), width=5.55 * inch, height=3.05 * inch),
            Paragraph(caption, styles["caption"]),
        ])

    w2_path = REPORTS_DIR / "W2_causal_attribution_report.pdf"
    w2 = SimpleDocTemplate(str(w2_path), pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=46, title="G9 因果归因与证据审计")
    story = [Paragraph("G9 因果归因与证据审计", styles["title"]), Paragraph("Executive Summary", styles["h1"])]
    bullets = [
        f"• 匹配后平衡通过，但面谈 ATT 为 {_pp(psm['estimate'])}，95% CI {_ci_pp(psm['ci95'])}，未发现可靠增益。",
        f"• AIPW 为 {_pp(aipw['estimate'])}，同样跨零。",
        f"• 延迟交付与评分的调整后差异为 {score['estimate']:+.3f} 分，置信区间跨零。",
        "• 数据不支持 RDD、多触点 Shapley 或渠道 ROI。",
    ]
    story += [Paragraph(item, styles["bullet"]) for item in bullets]
    story += [Paragraph("数据质量决定可用样本", styles["h1"]), Paragraph(f"5,000 条交付记录仅对应 {audit['row_counts']['unique_orders']:,} 个唯一订单；{audit['critical_conflict_orders']:,} 个订单存在关键字段冲突。主要交付分析只使用 {audit['usable_orders']:,} 个无关键冲突订单。", styles["body"])]
    story += image_block("propensity", "图 1. 倾向得分共同支持") + image_block("balance", "图 2. 匹配前后协变量平衡")
    story += [PageBreak(), Paragraph("跟进方式没有可靠的正向增益证据", styles["h1"])] + image_block("effects", "图 3. PSM 与 AIPW 效应估计") + image_block("sensitivity", "图 4. 隐藏偏差敏感性")
    data = [["分析", "估计", "95% CI", "证据等级"], ["PSM ATT", _pp(psm["estimate"]), _ci_pp(psm["ci95"]), psm["evidence_level"]], ["AIPW ATT", _pp(aipw["estimate"]), _ci_pp(aipw["ci95"]), aipw["evidence_level"]]]
    table = Table(data, colWidths=[1.2*inch,1.35*inch,2.2*inch,1.55*inch], repeatRows=1)
    table.setStyle(TableStyle([("FONTNAME",(0,0),(-1,-1),font_name),("FONTSIZE",(0,0),(-1,-1),8.5),("BACKGROUND",(0,0),(-1,0),colors.HexColor("#F2F4F7")),("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#CBD5E1")),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),6),("RIGHTPADDING",(0,0),(-1,-1),6),("TOPPADDING",(0,0),(-1,-1),6),("BOTTOMPADDING",(0,0),(-1,-1),6)])); story += [table]
    story += [PageBreak(), Paragraph("交付延迟只能报告调整后关联", styles["h1"]), Paragraph("评分结果跨零，不能支持显著下降。投诉结果即使通过 FDR，也只是调整后关联，不能替代延迟机制和完整时间戳。", styles["body"])] + image_block("delivery", "图 5. 延迟交付调整后关联")
    story += [Paragraph("RDD 可行性审计", styles["h1"]), Paragraph("没有连续分配变量和事先确定阈值。交付里程中位数不是延迟交付规则，因此停止生成 McCrary、带宽和安慰剂 RDD 结论。", styles["body"])]
    story += [Paragraph("渠道结果是关联性比较", styles["h1"]), Paragraph("每条线索只有一个渠道；没有多触点路径和渠道级花费。每百线索订单数不是 ROI。", styles["body"])] + image_block("channel", "图 6. 渠道转化率及 95% CI") + image_block("channel_standardized", "图 7. 标准化边际转化率")
    story += [Paragraph("Recommended Next Steps", styles["h1"])] + [Paragraph(item, styles["bullet"]) for item in ["• 补采下订、承诺交付和实际交付时间戳。", "• 运行跟进方式随机化试点。", "• 建立渠道级花费和多触点路径，并把证据等级、置信区间与看板建议同步。"]]
    w2.build(story, onFirstPage=_pdf_page, onLaterPages=_pdf_page)

    w5_path = REPORTS_DIR / "W5_sales_operations_white_paper.pdf"
    w5 = SimpleDocTemplate(str(w5_path), pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=46, title="G9 销售运营优化白皮书")
    white = [
        Paragraph("G9 销售运营优化白皮书", styles["title"]),
        Paragraph("按参考文档章节重构的证据审计版 · 54 张全中文图表", styles["body"]),
        Paragraph("沿用执行摘要、数据诊断、因果归因、预测预警、策略优化、实施路线图和方法附录结构；所有结论替换为当前可复现结果。", styles["body"]),
    ]
    color_legend = (
        f'<font color="{IMPROVE_RED}"><b>■ 优先改进</b></font>　'
        f'<font color="{WATCH_AMBER}"><b>■ 需要关注</b></font>　'
        f'<font color="{GOOD_GREEN}"><b>■ 表现较好</b></font>　'
        f'<font color="{NEUTRAL_BLUE}"><b>■ 描述性</b></font>'
    )
    white += [Paragraph(color_legend, styles["body"]), Paragraph("红橙绿按同图相对分位识别改进优先级，不代表业务目标、显著性或因果效应；蓝灰图只作描述。", styles["body"]), PageBreak()]
    specs = {item["key"]: item for item in WHITEPAPER_CHART_SEQUENCE}
    figure_no = 1
    for chapter_index, chapter in enumerate(WHITEPAPER_CHAPTER_STRUCTURE):
        if chapter_index:
            white.append(PageBreak())
        white.append(Paragraph(chapter["chapter"], styles["h1"]))
        white.append(Paragraph(f"<b>对应任务：</b>{chapter['task']} · {WHITEPAPER_TASK_SECTIONS[chapter['task']]['state']}", styles["body"]))
        for section in chapter["sections"]:
            white.append(Paragraph(section["title"], styles["h1"]))
            for paragraph in _whitepaper_section_copy(section["title"], analysis, audit):
                white.append(Paragraph(paragraph, styles["body"]))
            for key in section["keys"]:
                spec = specs[key]
                white.append(compact_image_block(key, f"图 {figure_no}. {spec['caption']}", spec["note"]))
                figure_no += 1
            if section["title"] == "5.4 综合策略推荐":
                white.append(PageBreak())
                strategy_rows = [["城市 / 渠道", "线索数", "转化率", "建议"]]
                for _, row in strategy.head(8).iterrows():
                    strategy_rows.append([f"{row['城市']} / {row['渠道']}", f"{row['线索数']:,}", f"{row['转化率']:.1%}", row["建议"]])
                strategy_table = Table(strategy_rows, colWidths=[1.35*inch, 0.7*inch, 0.8*inch, 3.45*inch], repeatRows=1)
                strategy_table.setStyle(TableStyle([("FONTNAME",(0,0),(-1,-1),font_name),("FONTSIZE",(0,0),(-1,-1),7.5),("BACKGROUND",(0,0),(-1,0),colors.HexColor("#F2F4F7")),("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#CBD5E1")),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)]))
                white.append(strategy_table)

    white += [PageBreak(), Paragraph("附录  方法论与技术栈", styles["h1"]), Paragraph("A.1 五周任务交付总览", styles["h1"])]
    task_rows = [["任务", "图表数", "完成状态", "主要输出"]]
    for task, meta in WHITEPAPER_TASK_SECTIONS.items():
        task_rows.append([task.split(" ", 1)[0], str(len(WHITEPAPER_TASK_CHART_KEYS[task])), meta["state"], meta["deliverable"]])
    task_table = Table(task_rows, colWidths=[0.65*inch, 0.6*inch, 1.75*inch, 3.25*inch], repeatRows=1)
    task_table.setStyle(TableStyle([("FONTNAME",(0,0),(-1,-1),font_name),("FONTSIZE",(0,0),(-1,-1),7.5),("BACKGROUND",(0,0),(-1,0),colors.HexColor("#F2F4F7")),("GRID",(0,0),(-1,-1),0.4,colors.HexColor("#CBD5E1")),("VALIGN",(0,0),(-1,-1),"MIDDLE"),("LEFTPADDING",(0,0),(-1,-1),5),("RIGHTPADDING",(0,0),(-1,-1),5),("TOPPADDING",(0,0),(-1,-1),5),("BOTTOMPADDING",(0,0),(-1,-1),5)]))
    white += [task_table, Paragraph("A.2 方法与复现", styles["h1"])]
    white += [Paragraph(item, styles["bullet"]) for item in ["• 原始 Excel 保持只读，清洗层、分析层、看板层均由派生数据生成。", "• 因果分析使用基线协变量匹配、共同支持、平衡诊断、增强逆概率加权和重采样。", "• 预测采用时间顺序验证，线上看板只展示聚合风险。", "• 无合法识别条件的方法保留可行性审计，不生成伪结论。"]]
    w5.build(white, onFirstPage=_pdf_page, onLaterPages=_pdf_page)
    return w2_path, w5_path


def write_supporting_reports(analysis: dict[str, Any], strategy: pd.DataFrame) -> None:
    model = analysis["model"]["metrics"]
    w3 = ["# W3 时间外预测验证", "", f"- 时间外 ROC AUC：{model['roc_auc']:.3f}", f"- PR AUC：{model['pr_auc']:.3f}", f"- Brier Score：{model['brier_score']:.3f}", "- 仅使用线索创建时可得字段。", "- 原伪造生存时间分析已删除。", ""]
    (REPORTS_DIR / "W3_model_validation.md").write_text("\n".join(w3), encoding="utf-8")
    columns = ["城市", "渠道", "线索数", "订单数", "转化率", "跟进覆盖率", "优先级分", "建议"]
    header = "| " + " | ".join(columns) + " |"
    divider = "|" + "|".join(["---"] * len(columns)) + "|"
    table_rows = []
    for _, row in strategy.head(15).iterrows():
        values = [row[c] for c in columns]
        values[4] = f"{float(values[4]):.1%}"; values[5] = f"{float(values[5]):.1%}"; values[6] = f"{float(values[6]):.1f}"
        table_rows.append("| " + " | ".join(str(v) for v in values) + " |")
    w4 = ["# W4 运营机会排序", "", "渠道级花费不可用，因此不计算 ROI/CPO。排序仅用于确定需要进一步诊断或试验的城市-渠道组合。", "", header, divider, *table_rows, ""]
    (REPORTS_DIR / "W4_strategy_report.md").write_text("\n".join(w4), encoding="utf-8")


def generate_all_reports(bundle: CleanBundle, analysis: dict[str, Any], charts: dict[str, pd.DataFrame]) -> dict[str, Path]:
    images = build_charts(charts, bundle, analysis)
    strategy = charts["strategy"]
    w2_docx = generate_w2_docx(analysis, bundle.audit, images)
    w5_docx, w5_md = generate_whitepaper(analysis, bundle.audit, images, strategy)
    w2_pdf, w5_pdf = generate_pdf_reports(analysis, bundle.audit, images, strategy)
    notebook = generate_notebook()
    write_supporting_reports(analysis, strategy)
    return {"w2_docx": w2_docx, "w2_pdf": w2_pdf, "w5_docx": w5_docx, "w5_pdf": w5_pdf, "w5_md": w5_md, "notebook": notebook, **images}
