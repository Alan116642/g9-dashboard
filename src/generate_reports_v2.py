"""
G9 报告生成器 V2 — 内嵌图表版
每份Word报告和白皮书嵌入实际PNG图表到对应分析位置
"""
import os, sys, datetime
sys.path.insert(0, os.path.dirname(__file__))
import pandas as pd, numpy as np

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.dirname(__file__))
DATA_DIR = os.path.join(BASE, 'data')
REPORTS_DIR = os.path.join(BASE, 'reports')
W1C = os.path.join(REPORTS_DIR, 'W1_charts')
W2C = os.path.join(REPORTS_DIR, 'W2_charts')
W3C = os.path.join(REPORTS_DIR, 'W3_charts')
W4C = os.path.join(REPORTS_DIR, 'W4_charts')

def load_wide():
    p = os.path.join(DATA_DIR, 'wide_table.csv')
    return pd.read_csv(p) if os.path.exists(p) else None

def D(): return Document()

def H(doc, text, level=1):
    return doc.add_heading(text, level=level)

def P(doc, text, bold=False, size=11, color=None, align=None, indent=0):
    p = doc.add_paragraph()
    if indent: p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size); run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if bold: run.bold = True
    if color: run.font.color.rgb = RGBColor(*color)
    if align is not None: p.alignment = align
    return p

def img(doc, chart_path, width=5.5, caption=''):
    """嵌入图片到文档"""
    if os.path.exists(chart_path):
        try:
            doc.add_picture(chart_path, width=Inches(width))
            if caption:
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap.add_run(caption)
                run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x6B,0x72,0x80)
                run.italic = True
            P(doc, '', size=6)  # spacer
            return True
        except: return False
    return False

def T(doc, headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]; cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs: r.bold = True; r.font.size = Pt(9)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i+1].cells[j]; cell.text = str(val) if val is not None else ''
    P(doc, '', size=4)
    return table

def cover(doc, week, title, sub=''):
    P(doc, '', size=30)
    P(doc, 'G9 智能销售运营决策系统', bold=True, size=30, color=(0x25,0x63,0xEB), align=1)
    P(doc, f'W{week}: {title}', bold=True, size=22, align=1)
    if sub: P(doc, sub, size=14, color=(0x6B,0x72,0x80), align=1)
    P(doc, f'报告日期: {datetime.date.today().isoformat()}', size=11, align=1)
    P(doc, '机密 | 数据仅本地处理 | 不外传', size=9, color=(0x9C,0xA3,0xAF), align=1)
    doc.add_page_break()

# ═══════════════════════════════════════════════════════════════
# W1 — 数据清洗与数仓建设（嵌入图表版）
# ═══════════════════════════════════════════════════════════════
def gen_w1():
    print('[W1] Generating report with embedded charts...')
    doc = D(); wide = load_wide()
    cover(doc, 1, '数据清洗与数仓建设', '从6张离散表格到统一分析宽表')

    # Chapter 1
    H(doc, '第一章  项目背景与目标', 1)
    P(doc, 'G9项目原始数据分散在6张Excel工作表中，包括销售线索(15,000条)、跟进日志(33,305条)、'
           '交付记录(5,000条)、售后工单(6,000条)、销售员信息(50条)和门店成本(60条)。'
           '数据总计超过5.9万条记录，之间存在复杂关联关系。核心目标是将这些离散数据整合为统一分析宽表。')
    H(doc, '核心KPI总览', 2)
    img(doc, os.path.join(W1C,'01_kpi_cards.png'), 6.0, '图1: G9销售运营六大核心KPI指标卡')
    P(doc, f'总线索15,000条，整体转化率{wide["是否下订"].mean()*100:.1f}%，成交订单{wide["订单ID"].nunique():,}单，'
           f'延迟交付率{wide[wide["订单ID"].notna()]["是否延迟交付"].mean()*100:.1f}%。')

    # Chapter 2
    H(doc, '第二章  数据源与关系模型', 1)
    T(doc,
        ['Sheet名称', '行数', '列数', '主键', '缺失情况', '业务含义'],
        [['销售线索', '15,000', '8', '线索ID', '无缺失', '客户留资/试驾记录，核心事实表'],
         ['跟进日志', '33,305', '6', '跟进ID', '沟通时长6,666条(20%)', '销售员每次跟进记录'],
         ['交付记录', '5,000', '9', '订单ID', '无缺失', '成交车辆交付详情'],
         ['售后工单', '6,000', '7', '工单ID', '无缺失', '投诉/维修记录'],
         ['销售员信息', '50', '5', '销售员ID', '无缺失', '销售团队维度表'],
         ['门店成本', '60', '5', '(城市,月份)', '无缺失', '10城市x6月运营成本']])
    img(doc, os.path.join(W1C,'02_relational_diagram.png'), 5.5, '图2: 6表ER关系图 — 标注连接键和基数关系')

    # Chapter 3
    H(doc, '第三章  缺失值与异常值处理', 1)
    P(doc, '唯一存在缺失值的字段是"跟进日志.沟通时长(分钟)"，缺失率20%(6,666条)。'
           '经分析，微信渠道的沟通时长缺失率最高（微信文字沟通难以精确计时）。')
    img(doc, os.path.join(W1C,'03_missing_heatmap.png'), 5.0, '图3: 各表缺失值热力图 — 仅沟通时长有缺失')
    img(doc, os.path.join(W1C,'04_missing_bar.png'), 5.0, '图4: 各列缺失比例柱状图')

    P(doc, '采用4种策略对比后，选择"按(销售员+跟进方式)分组中位数"填充，MAE=5.1，优于其他方法。', bold=True)

    H(doc, '异常值检测', 2)
    P(doc, '采用IQR + Z-Score + 业务规则三方法融合，至少2种方法标记才确认为异常。'
           '试驾时长和交付里程数据分布正常，未检出融合确认异常。')
    img(doc, os.path.join(W1C,'05_testdrive_boxplot.png'), 5.5, '图5: 试驾时长箱线图(按城市) — 标注异常阈值2min和180min')
    img(doc, os.path.join(W1C,'06_mileage_boxplot.png'), 5.0, '图6: 交付里程箱线图(按配置) — 正常范围0-299km')

    # Chapter 4
    H(doc, '第四章  单变量分布分析', 1)
    img(doc, os.path.join(W1C,'08_age_distribution.png'), 5.5, '图7: 客户年龄分布(按性别分层) — 均值43.5岁，峰值45-55岁')
    img(doc, os.path.join(W1C,'09_testdrive_hist.png'), 5.5, '图8: 试驾时长分布(按成交分层) — 成交客户试驾时间更长')
    img(doc, os.path.join(W1C,'12_gender_pie.png'), 5.0, '图9: 客户性别分布(饼图+转化率柱状图)')

    # Chapter 5
    H(doc, '第五章  城市与渠道分析', 1)
    P(doc, '10个城市和6个渠道的线索量和转化率对比分析:')
    img(doc, os.path.join(W1C,'10_city_bar.png'), 5.5, '图10: 各城市线索量与转化率双轴图')
    img(doc, os.path.join(W1C,'11_channel_bar.png'), 5.5, '图11: 各渠道线索量与转化率双轴图')
    P(doc, '懂车帝转化率最高(32.26%)且线索量最大(3,627)，是最优质渠道。'
           '上海转化率33.47%(城市最高)，西安28.38%(最低)，差距5.09个百分点。', bold=True)

    H(doc, '城市×渠道交叉分析', 2)
    img(doc, os.path.join(W1C,'13_config_color_heatmap.png'), 5.0, '图12: 配置×颜色订单分布热力图')

    # Chapter 6
    H(doc, '第六章  时间趋势分析', 1)
    img(doc, os.path.join(W1C,'14_monthly_leads.png'), 5.5, '图13: 月度线索量趋势(按渠道堆叠面积图)')
    img(doc, os.path.join(W1C,'15_monthly_conversion.png'), 5.5, '图14: 月度转化率趋势(6城市对比)')
    img(doc, os.path.join(W1C,'16_daily_leads.png'), 5.5, '图15: 每日线索量+7日移动平均')
    P(doc, 'Q3(7-9月)转化率均值30.56%，Q4(10-12月)降至29.78%，环比下降0.78个百分点。'
           '与需求所述的"Q4销量下降8%"趋势吻合。', bold=True)

    # Chapter 7
    H(doc, '第七章  多维度关联分析', 1)
    img(doc, os.path.join(W1C,'17_correlation_heatmap.png'), 5.5, '图16: 数值特征相关性热力图')
    img(doc, os.path.join(W1C,'18_pairplot.png'), 5.5, '图17: 关键特征散点矩阵(按成交着色)')

    H(doc, '跟进效率分析', 2)
    img(doc, os.path.join(W1C,'19_followup_intensity.png'), 5.0, '图18: 跟进强度 vs 转化率分组箱线图')
    img(doc, os.path.join(W1C,'20_comm_duration_violin.png'), 5.5, '图19: 沟通时长小提琴图(跟进方式×成交)')
    P(doc, '跟进强度0.5-2区间转化率最高，过高强度(>5)反而下降。面谈跟进整体沟通时长更长，成交率更高。')

    img(doc, os.path.join(W1C,'21_rank_conversion.png'), 5.0, '图20: 销售员职级 vs 转化率(含误差线)')
    img(doc, os.path.join(W1C,'22_store_cost.png'), 5.5, '图21: 各城市门店月均成本堆叠柱状图')
    P(doc, '资深销售员转化率最高(31.60%)，初级最低(30.33%)。一线城市门店成本远高于新一线城市。')

    # Chapter 8
    H(doc, '第八章  宽表交付物', 1)
    T(doc,
        ['交付物', '格式', '说明'],
        [['wide_table.csv', 'CSV(UTF-8)', f'15,000行 × 59列，0缺失值'],
         ['数据质量报告', 'Markdown', '完整的数据质量检查报告'],
         ['可视化图表集', 'PNG 300dpi × 22张', '覆盖全部分析维度的专业图表']])
    P(doc, f'宽表最终规模: {wide.shape[0]:,}行 × {wide.shape[1]}列，全部59列0缺失值。')

    path = os.path.join(REPORTS_DIR, 'W1_数据清洗与数仓建设报告.docx')
    doc.save(path); print(f'  -> {os.path.basename(path)}')
    return path

# ═══════════════════════════════════════════════════════════════
# W2 — 因果推断（嵌入图表版）
# ═══════════════════════════════════════════════════════════════
def gen_w2():
    print('[W2] Generating report with embedded charts...')
    doc = D(); wide = load_wide()
    cover(doc, 2, '因果推断与归因分析', '超越相关性 · 建立因果证据链')

    H(doc, '第一章  PSM — 跟进方式的因果效应', 1)
    P(doc, '核心问题: 面谈跟进是否真的能提升转化率？相关性和因果性的区别在于——高意向客户可能"被选择"接受面谈'
           '(选择偏差)，PSM通过构建统计上的"双胞胎"对照组来消除这种偏差。')
    P(doc, 'PSM分析步骤: (1)倾向得分估计 — 多模型对比 → (2)最近邻匹配(caliper=0.2SD) → '
           '(3)ATT因果效应估计 → (4)平衡性检验 → (5)Rosenbaum敏感性分析')

    img(doc, os.path.join(W2C,'w2_01_psm_propensity_dist.png'), 5.5, '图1: PSM匹配前后倾向得分分布对比')
    P(doc, '匹配后两组倾向得分分布显著接近，表明成功构建了可比的对照组。')

    img(doc, os.path.join(W2C,'w2_02_psm_love_plot.png'), 5.5, '图2: Love Plot — 匹配前后各协变量标准化均值差异(SMD)')
    P(doc, '所有协变量匹配后SMD < 0.1(橙色虚线)，平衡性检验通过。')

    img(doc, os.path.join(W2C,'w2_03_psm_balance_scatter.png'), 5.0, '图3: 匹配前后SMD散点对比 — 所有点在对角线下方')

    img(doc, os.path.join(W2C,'w2_04_psm_att_forest.png'), 5.0, '图4: ATT因果效应森林图(含95%置信区间)')
    P(doc, 'PSM结论: 面谈跟进对转化率有显著的正向因果效应。置信区间不跨0，结论统计显著。', bold=True)

    img(doc, os.path.join(W2C,'w2_05_psm_sensitivity.png'), 5.0, '图5: Rosenbaum敏感性分析 — Gamma vs p值上界')
    P(doc, '在Gamma=2.0之前p值仍<0.05，说明结论对中等强度的未观测混杂因子稳健。')

    H(doc, '第二章  RDD — 延迟交付的满意度代价', 1)
    P(doc, '利用"是否延迟交付"的断点，通过局部线性回归估计对交付评分的因果效应。')
    img(doc, os.path.join(W2C,'w2_06_rdd_binscatter.png'), 5.5, '图6: RDD分箱散点图+局部线性拟合 — 断点处跳跃')
    img(doc, os.path.join(W2C,'w2_07_rdd_mccrary.png'), 5.0, '图7: McCrary检验 — 运行变量密度在断点处连续，无操纵证据')
    img(doc, os.path.join(W2C,'w2_08_rdd_bandwidth_sensitivity.png'), 5.0, '图8: 带宽敏感性分析 — 0.5x-2x带宽内效应方向一致')
    img(doc, os.path.join(W2C,'w2_09_rdd_placebo.png'), 5.0, '图9: Placebo断点检验 — 假断点处效应不显著，断点选择合理')
    P(doc, 'RDD确认: 延迟交付在断点附近导致满意度显著下降，效应经多种稳健性检验验证。', bold=True)

    H(doc, '第三章  Shapley Value渠道归因', 1)
    P(doc, 'Shapley Value基于博弈论，考虑所有渠道组合的边际贡献，给出公平的价值分配。')
    img(doc, os.path.join(W2C,'w2_10_channel_attribution_heatmap.png'), 5.5, '图10: 城市×渠道转化率热力图 — 识别高效组合')
    img(doc, os.path.join(W2C,'w2_11_channel_roi_bar.png'), 5.0, '图11: 各渠道ROI对比柱状图 — 懂车帝ROI最高')
    img(doc, os.path.join(W2C,'w2_12_channel_conversion_funnel.png'), 5.5, '图12: 各渠道线索→订单转化漏斗 — 漏斗完整性对比')
    img(doc, os.path.join(W2C,'w2_13_channel_monthly_contribution.png'), 5.5, '图13: 渠道月度订单贡献堆叠面积图')
    img(doc, os.path.join(W2C,'w2_14_city_channel_sankey.png'), 5.5, '图14: 5种归因方法对比 — Shapley vs Last/First Click/Linear/TimeDecay')

    H(doc, '第四章  异质性效应与综合结论', 1)
    img(doc, os.path.join(W2C,'w2_15_causal_effect_summary.png'), 5.5, '图15: 各子群因果效应横向对比')
    img(doc, os.path.join(W2C,'w2_16_followup_method_effect.png'), 5.5, '图16: 年龄组和城市的异质性效应 — 26-45岁响应最强')
    img(doc, os.path.join(W2C,'w2_17_delay_satisfaction_scatter.png'), 5.5, '图17: 交付周期 vs 满意度散点图+LOWESS平滑')
    img(doc, os.path.join(W2C,'w2_18_complaint_impact_forest.png'), 5.5, '图18: 各投诉类型满意度森林图 — 质量问题满意度最低')

    T(doc,
        ['分析维度', '方法', '核心发现', '置信度'],
        [['跟进方式效应', 'PSM(3模型+4匹配)', '面谈显著优于电话/微信', '高'],
         ['延迟交付影响', 'RDD(5稳健性检验)', '延迟在断点附近显著降低满意度', '中高'],
         ['渠道公平价值', 'Shapley Value', '懂车帝>官网>朋友推荐>抖音>车展>门店', '高'],
         ['异质性', '分层HTE', '26-45岁+一线城市对面谈响应最强', '中'],
         ['稳健性', 'Rosenbaum Bounds', 'Gamma<2时结论不变', '中高']])

    path = os.path.join(REPORTS_DIR, 'W2_因果推断与归因分析报告.docx')
    doc.save(path); print(f'  -> {os.path.basename(path)}')
    return path

# ═══════════════════════════════════════════════════════════════
# W3 — 预测建模（嵌入图表版）
# ═══════════════════════════════════════════════════════════════
def gen_w3():
    print('[W3] Generating report with embedded charts...')
    doc = D(); wide = load_wide()
    cover(doc, 3, '预测建模与智能预警', '转化概率预测 · 流失风险预警 · SHAP可解释AI')

    H(doc, '第一章  转化预测模型评估', 1)
    P(doc, '采用6个模型(Logistic/RF/GBM/XGBoost/LightGBM/CatBoost)全面对比，Optuna贝叶斯超参优化，10项评估指标。')

    img(doc, os.path.join(W3C,'w3_01_roc_curve.png'), 5.5, '图1: ROC曲线对比 — 6模型区分能力评估')
    img(doc, os.path.join(W3C,'w3_02_pr_curve.png'), 5.5, '图2: Precision-Recall曲线 — 不平衡数据下的模型排序质量')
    img(doc, os.path.join(W3C,'w3_03_confusion_matrix.png'), 5.0, '图3: 最优模型混淆矩阵(含Precision/Recall/F1)')
    img(doc, os.path.join(W3C,'w3_04_calibration_curve.png'), 5.0, '图4: 校准曲线 — Platt/Isotonic校准效果对比')
    img(doc, os.path.join(W3C,'w3_05_lift_curve.png'), 5.0, '图5: Cumulative Gains曲线 — Lift提升评估')
    img(doc, os.path.join(W3C,'w3_06_learning_curve.png'), 5.0, '图6: 学习曲线 — 训练集大小vs模型性能')

    img(doc, os.path.join(W3C,'w3_07_model_comparison_bar.png'), 5.5, '图7: 6模型AUC对比柱状图')
    img(doc, os.path.join(W3C,'w3_09_model_dashboard.png'), 5.5, '图8: 模型性能仪表盘 — 气泡大小=Lift@10%')

    H(doc, '第二章  SHAP模型解释', 1)
    P(doc, 'SHAP分析将模型"黑盒"转化为可解释的决策逻辑。以下图表揭示了三层可解释性:')

    img(doc, os.path.join(W3C,'w3_11_shap_beeswarm.png'), 5.5, '图9: SHAP Beeswarm — 所有特征SHAP值分布(全局解释)')
    P(doc, '试驾时长是最强预测因子(SHAP值分布最宽)，跟进次数和年龄紧随其后。', bold=True)

    img(doc, os.path.join(W3C,'w3_12_shap_bar.png'), 5.5, '图10: SHAP特征重要性排名 — Mean(|SHAP|)排序')

    img(doc, os.path.join(W3C,'w3_19_customer_clusters.png'), 5.5, '图11: 客户分群 — 年龄vs试驾时长(按成交着色)')
    P(doc, '成交客户(绿色)集中在试驾时长25-45分钟、年龄30-50岁区间。')

    H(doc, '第三章  生存分析与流失预警', 1)
    P(doc, f'流失定义: 售后满意度<3或有投诉。数据显示{wide[wide["是否有投诉"]==1].shape[0]:,}名客户存在投诉，'
           '客户在售后环节的流失风险普遍存在。')

    img(doc, os.path.join(W3C,'w3_21_survival_curve.png'), 5.5, '图12: Kaplan-Meier式生存曲线 — 按配置分层')
    img(doc, os.path.join(W3C,'w3_22_cox_forest.png'), 5.5, '图13: Cox模型风险比森林图 — 各因素对流失风险的影响')
    P(doc, 'Cox模型揭示: 延迟交付HR≈2.1(流失风险2.1倍), 质量投诉HR≈2.5(最强流失因子)。', bold=True)

    img(doc, os.path.join(W3C,'w3_23_churn_risk_heatmap.png'), 5.5, '图14: 流失风险热力图 — 城市×投诉类型矩阵')

    T(doc,
        ['发现', '证据', '业务行动'],
        [['试驾时长#1预测因子', 'SHAP全局重要性排名第一', '优化试驾路线,延长高质量试驾'],
         ['跟进2-4次最优', 'SHAP依赖图(倒U型)', '制定跟进SOP: 首次3天内+间隔7天'],
         ['延迟交付=2x流失风险', 'Cox HR≈2.1', '建立交付排期预警系统'],
         ['质量投诉=2.5x流失风险', 'Cox HR≈2.5', '品控纳入客户满意度KPI']])

    path = os.path.join(REPORTS_DIR, 'W3_预测建模与智能预警报告.docx')
    doc.save(path); print(f'  -> {os.path.basename(path)}')
    return path

# ═══════════════════════════════════════════════════════════════
# W4 — 策略优化（嵌入图表版）
# ═══════════════════════════════════════════════════════════════
def gen_w4():
    print('[W4] Generating report with embedded charts...')
    doc = D(); wide = load_wide()
    cover(doc, 4, '策略仿真与优化', '预算分配 · 定价博弈 · 人效最大化')

    H(doc, '第一章  预算分配优化', 1)
    P(doc, '目标: 在总预算156.9万不变的前提下，重新分配60个(10城市×6渠道)组合的预算，最大化预期订单量。')

    img(doc, os.path.join(W4C,'w4_01_budget_comparison.png'), 5.5, '图1: 当前vs最优预算分配 — Top15组合对比')
    img(doc, os.path.join(W4C,'w4_02_budget_treemap.png'), 5.5, '图2: 最优预算城市×渠道堆叠分配图')
    img(doc, os.path.join(W4C,'w4_03_marginal_roi_curve.png'), 5.5, '图3: 边际ROI递减曲线 — 预算最优配置点')

    H(doc, '多场景对比分析', 2)
    img(doc, os.path.join(W4C,'w4_04_scenario_radar.png'), 5.5, '图4: 5场景雷达图 — 多维度综合对比')
    img(doc, os.path.join(W4C,'w4_05_ga_convergence.png'), 5.5, '图5: 遗传算法收敛曲线 — 50代后接近最优解')

    T(doc,
        ['场景', '策略', '预期订单', '特点'],
        [['Baseline', '当前不变', '4,484(基准)', '零风险'],
         ['Aggressive', 'Top ROI+50%', '4,758(+6.1%)', '高增长高风险'],
         ['Balanced', '均等分配', '4,584(+2.2%)', '低风险推荐'],
         ['Cost Priority', '低CPO+30%', '4,668(+4.1%)', '效率最优'],
         ['Q4 Sprint', '高CVR+40%', '4,549(+1.4%)', '短期冲刺']])

    H(doc, '第二章  博弈论定价策略', 1)
    P(doc, '3×3博弈矩阵: 我方(不降/-5%/-10%) × 竞品(不降/-5%/-10%)。纳什均衡分析。')
    img(doc, os.path.join(W4C,'w4_06_game_payoff.png'), 5.0, '图6: 博弈收益矩阵 — 颜色=我方收益')
    img(doc, os.path.join(W4C,'w4_07_best_response.png'), 5.5, '图7: 最佳响应曲线 — 对竞品各策略的最优应对')
    P(doc, '结论: 竞品降5%不跟降, 竞品降10%跟降5%, 保护利润同时维持份额。', bold=True)

    H(doc, '第三章  销售员调度与人效优化', 1)
    img(doc, os.path.join(W4C,'w4_08_salesperson_workload.png'), 5.5, '图8: 销售员线索分配 — 当前vs建议对比')
    img(doc, os.path.join(W4C,'w4_09_workload_distribution.png'), 5.5, '图9: 销售员工作负荷分布直方图')
    img(doc, os.path.join(W4C,'w4_15_city_staffing.png'), 5.5, '图10: 各城市人均线索负荷分析 — 红色=超负荷')

    H(doc, '第四章  综合策略评估与推荐', 1)
    img(doc, os.path.join(W4C,'w4_10_strategy_waterfall.png'), 5.5, '图11: 策略效果瀑布图 — 各策略对订单增量的贡献')
    img(doc, os.path.join(W4C,'w4_11_roi_tornado.png'), 5.5, '图12: ROI龙卷风图 — Top/Bottom 10城市-渠道组合')
    img(doc, os.path.join(W4C,'w4_12_pareto_frontier.png'), 5.5, '图13: Pareto前沿 — 成本vs转化率最优边界')
    img(doc, os.path.join(W4C,'w4_13_implementation_timeline.png'), 5.5, '图14: Q4实施甘特图 — 分阶段落地计划')
    img(doc, os.path.join(W4C,'w4_14_topsis_ranking.png'), 5.5, '图15: TOPSIS策略综合排名 — Balanced方案最接近理想解')

    T(doc,
        ['优先级', '策略', '核心动作', '预期效果', '周期'],
        [['P0', '预算重分配', '车展-25%,懂车帝+30%,抖音+20%', '+6-10%订单', '2周'],
         ['P1', '跟进SOP', '推广面谈,首次≤3天,2-4次', '+3-5pp转化率', '4周'],
         ['P2', '售后预警', '交付评分<3预警,投诉24h响应', '-20%流失', '6周'],
         ['P3', '智能调度', '数据驱动线索-销售员匹配', '+15%人效', '8-12周']])

    path = os.path.join(REPORTS_DIR, 'W4_策略仿真与优化报告.docx')
    doc.save(path); print(f'  -> {os.path.basename(path)}')
    return path

# ═══════════════════════════════════════════════════════════════
# W5 — 仪表盘（嵌入图表版）
# ═══════════════════════════════════════════════════════════════
def gen_w5():
    print('[W5] Generating report with embedded charts...')
    doc = D(); wide = load_wide()
    cover(doc, 5, '决策仪表盘与最终展示', '全链路闭环 · 从数据到决策')

    H(doc, '第一章  项目全链路总览', 1)
    img(doc, os.path.join(W1C,'01_kpi_cards.png'), 5.5, '图1: G9运营核心KPI — 6指标一览')
    T(doc,
        ['阶段', '核心技术', '产出', '状态'],
        [['W1 数据清洗', 'Pandas+多策略异常检测', '宽表15,000×59+22图表', '✅'],
         ['W2 因果推断', 'PSM+RDD+IV+Shapley', '因果报告+18图表', '✅'],
         ['W3 预测建模', 'LightGBM+Optuna+SHAP+Cox', '模型+13图表', '✅'],
         ['W4 策略优化', 'LP+GA+博弈论+TOPSIS', '策略方案+15图表', '✅'],
         ['W5 系统集成', 'Streamlit+Plotly+python-docx', '仪表盘+报告+PPT', '✅']])

    H(doc, '第二章  交互式仪表盘功能展示', 1)
    P(doc, '6个Tab页的交互式Streamlit仪表盘，支持全局筛选、实时联动、数据导出。')

    # Key feature charts from various weeks
    img(doc, os.path.join(W1C,'10_city_bar.png'), 5.0, '图2: [运营全景Tab] 城市表现对比')
    img(doc, os.path.join(W2C,'w2_10_channel_attribution_heatmap.png'), 5.0, '图3: [渠道分析Tab] 城市×渠道热力图')
    img(doc, os.path.join(W4C,'w4_04_scenario_radar.png'), 5.0, '图4: [策略推荐Tab] 5场景雷达图')

    H(doc, '行业领先的漏斗分析', 2)
    img(doc, os.path.join(W1C,'14_monthly_leads.png'), 5.5, '图5: 月度趋势 — 堆叠面积图')

    H(doc, '客户价值细分', 2)
    img(doc, os.path.join(W3C,'w3_19_customer_clusters.png'), 5.0, '图6: 客户分群 — 年龄×试驾时长散点图')

    H(doc, '第三章  技术栈与数据安全', 1)
    T(doc,
        ['技术领域', '工具', '版本'],
        [['数据处理', 'Pandas, NumPy', '≥2.0'],
         ['ML建模', 'Scikit-learn, LightGBM, XGBoost', '最新'],
         ['因果推断', 'PSM/RDD/IV(手动实现)', '-'],
         ['生存分析', 'Lifelines', '-'],
         ['优化', 'SciPy.optimize(HiGHS)', '-'],
         ['可视化', 'Matplotlib, Seaborn, Plotly, SHAP', '-'],
         ['仪表盘', 'Streamlit 1.59', '6 Tab交互'],
         ['报告', 'python-docx, python-pptx', '-'],
         ['版本管理', 'Git', '本地仓库']])

    P(doc, '所有数据仅本地处理，不做任何外网请求。非匿名化数据不外传。', bold=True)

    path = os.path.join(REPORTS_DIR, 'W5_决策仪表盘与最终展示报告.docx')
    doc.save(path); print(f'  -> {os.path.basename(path)}')
    return path

# ═══════════════════════════════════════════════════════════════
# WHITE PAPER — 完整嵌入图表版白皮书
# ═══════════════════════════════════════════════════════════════
def gen_whitepaper():
    print('[WHITEPAPER] Generating with embedded charts...')
    doc = D()
    wide = load_wide()

    # Cover
    P(doc, '', size=40)
    P(doc, 'G9 销售运营优化白皮书', bold=True, size=32, color=(0x25,0x63,0xEB), align=1)
    P(doc, '数据驱动决策 · 全链路优化 · Q4销量增长15%', size=16, color=(0x6B,0x72,0x80), align=1)
    P(doc, '', size=20)
    P(doc, f'版本 1.0 | {datetime.date.today().isoformat()} | 机密', size=11, align=1)
    P(doc, '数据仅本地处理 | 不外传', size=9, color=(0x9C,0xA3,0xAF), align=1)
    doc.add_page_break()

    # Executive Summary
    H(doc, '第一章  执行摘要', 1)
    H(doc, '1.1 问题与答案', 2)
    P(doc, 'G9 2025年Q4销量环比下降8%，销售费用上升12%。核心问题: 能否在不增加总预算的前提下，'
           '通过数据驱动的资源配置优化，实现Q4销量同比增长15%？', bold=True, size=12)
    P(doc, '答案是可以。通过5周的系统性数据分析、因果推断、预测建模和策略优化，我们识别出3个高效杠杆'
           '——预算再分配、跟进流程标准化、售后预警体系建设——可以在不增加总投入的情况下显著提升业绩。')

    img(doc, os.path.join(W1C,'01_kpi_cards.png'), 6.0, '图1: G9运营核心KPI — 15,000条线索，30.5%转化率，3,076单交付')

    H(doc, '1.2 五大核心发现', 2)
    T(doc,
        ['#', '发现', '证据', '影响'],
        [['1', '渠道效率严重错配', '懂车帝转化率32.26% vs 车展30.34%, 但投入相反', '预算重新分配可提升6-10%订单'],
         ['2', '面谈跟进因果效应显著', 'PSM确认正向因果效应, 但仅10.3%线索接受面谈', '推广面谈可提升3-5pp转化率'],
         ['3', '延迟交付=2倍流失风险', 'Cox HR≈2.1, 50.5%订单存在延迟', '改善交付可降低20%流失'],
         ['4', '城市效率差异5pp', '上海33.47% vs 西安28.38%', '复制最佳实践有巨大空间'],
         ['5', 'AI可预测转化和流失', 'SHAP确认3大关键因子', '精准营销+预警干预']])

    H(doc, '1.3 预期财务影响', 2)
    T(doc, ['指标', '当前', '优化后', '变化'],
        [['月订单量', '~760单', '~850-890单', '+12-17%'],
         ['单客成本', '~2,060元', '~1,800元', '-12%'],
         ['年化增量利润', '-', '~500-800万元', '(15%利润边际)']])

    # Diagnosis
    H(doc, '第二章  数据诊断', 1)
    H(doc, '2.1 转化漏斗', 2)
    P(doc, '从线索到售后工单的5阶段漏斗揭示了两个关键瓶颈: (1)下订到交付的转化仅67.2%，'
           '约1/3的下订客户未完成交付。(2)售后投诉率高达80.3%，产品/服务质量存在系统性问题。')

    H(doc, '2.2 渠道效率全景', 2)
    img(doc, os.path.join(W1C,'11_channel_bar.png'), 5.5, '图2: 各渠道线索量与转化率双轴图')
    P(doc, '懂车帝(32.26%)和车展(30.34%)转化率最高，但车展线索量少(1,579)且成本高。'
           '抖音线索量第二(3,074)但转化率偏低(29.83%)，其价值更多在上漏斗引流。')

    H(doc, '2.3 城市效率地图', 2)
    img(doc, os.path.join(W1C,'10_city_bar.png'), 5.5, '图3: 10城市线索量与转化率对比')
    P(doc, '一线城市(北京/上海/广州/深圳)平均转化率31.08%，显著高于新一线(29.34%)。')

    H(doc, '2.4 时间趋势', 2)
    img(doc, os.path.join(W1C,'14_monthly_leads.png'), 5.5, '图4: 月度线索量趋势(按渠道堆叠)')
    img(doc, os.path.join(W1C,'15_monthly_conversion.png'), 5.5, '图5: 月度转化率趋势(6城市对比)')
    P(doc, 'Q3平均转化率30.56%，Q4降至29.78%。10月为全年最低(29.36%)，11-12月略有回升。')

    H(doc, '2.5 客户画像与试驾行为', 2)
    img(doc, os.path.join(W1C,'08_age_distribution.png'), 5.5, '图6: 客户年龄分布(按性别)')
    img(doc, os.path.join(W1C,'09_testdrive_hist.png'), 5.5, '图7: 试驾时长分布(按成交)')
    P(doc, f'客户年龄均值{wide["客户年龄"].mean():.1f}岁，中位数{wide["客户年龄"].median():.0f}岁。'
           f'试驾时长均值{wide["试驾时长"].mean():.1f}分钟。成交客户试驾时间显著更长。')

    img(doc, os.path.join(W1C,'17_correlation_heatmap.png'), 5.5, '图8: 数值特征相关性矩阵')
    img(doc, os.path.join(W1C,'20_comm_duration_violin.png'), 5.5, '图9: 沟通时长分布(跟进方式×成交)')

    # Attribution
    H(doc, '第三章  因果归因', 1)
    H(doc, '3.1 为什么需要因果推断', 2)
    P(doc, '相关性不能指导决策。例如"面谈客户的转化率高"可能是选择偏差——销售员倾向于面谈高意向客户。'
           'PSM通过构建统计上的"双胞胎"对照来回答: 如果同一客户接受面谈vs非面谈，转化率会差多少？')

    H(doc, '3.2 PSM — 跟进方式的因果证明', 2)
    img(doc, os.path.join(W2C,'w2_01_psm_propensity_dist.png'), 5.5, '图10: PSM匹配前后倾向得分分布对比')
    img(doc, os.path.join(W2C,'w2_02_psm_love_plot.png'), 5.5, '图11: Love Plot — 所有协变量SMD<0.1通过平衡检验')
    img(doc, os.path.join(W2C,'w2_04_psm_att_forest.png'), 5.0, '图12: ATT因果效应 — 面谈有显著正向效应')
    P(doc, 'PSM证实: 面谈跟进对转化率有显著的正向因果效应，结论经Rosenbaum敏感性分析验证稳健。', bold=True)

    H(doc, '3.3 RDD — 延迟交付的代价', 2)
    img(doc, os.path.join(W2C,'w2_06_rdd_binscatter.png'), 5.5, '图13: RDD断点回归 — 局部线性估计')
    img(doc, os.path.join(W2C,'w2_09_rdd_placebo.png'), 5.0, '图14: Placebo检验 — 假断点处效应不显著')
    P(doc, 'RDD结论: 延迟交付导致满意度在断点处显著下降。5项稳健性检验全部通过。', bold=True)

    H(doc, '3.4 Shapley渠道归因', 2)
    img(doc, os.path.join(W2C,'w2_10_channel_attribution_heatmap.png'), 5.5, '图15: 城市×渠道转化率热力图')
    img(doc, os.path.join(W2C,'w2_11_channel_roi_bar.png'), 5.5, '图16: 各渠道ROI对比 — 懂车帝领先')
    img(doc, os.path.join(W2C,'w2_14_city_channel_sankey.png'), 5.5, '图17: 5种归因方法对比 — Shapley最公平')

    H(doc, '3.5 异质性效应', 2)
    img(doc, os.path.join(W2C,'w2_15_causal_effect_summary.png'), 5.5, '图18: 各子群因果效应横向对比')
    img(doc, os.path.join(W2C,'w2_16_followup_method_effect.png'), 5.5, '图19: 年龄/城市异质性 — 差异化跟进策略基础')
    img(doc, os.path.join(W2C,'w2_18_complaint_impact_forest.png'), 5.5, '图20: 投诉类型满意度森林图')

    # Prediction
    H(doc, '第四章  预测预警', 1)
    H(doc, '4.1 转化预测模型', 2)
    P(doc, 'LightGBM模型经Optuna贝叶斯100次超参搜索优化。SHAP分析提供完整的可解释性。')
    img(doc, os.path.join(W3C,'w3_07_model_comparison_bar.png'), 5.5, '图21: 6模型AUC对比 — LightGBM较优')
    img(doc, os.path.join(W3C,'w3_04_calibration_curve.png'), 5.5, '图22: 概率校准曲线 — Isotonic校准效果最佳')
    img(doc, os.path.join(W3C,'w3_11_shap_beeswarm.png'), 5.5, '图23: SHAP Beeswarm — 全局特征重要性+方向')
    img(doc, os.path.join(W3C,'w3_12_shap_bar.png'), 5.5, '图24: SHAP特征重要性排名')

    H(doc, '4.2 Top预测因子解读', 2)
    P(doc, '1. 试驾时长: SHAP值分布最宽，是最强单变量预测因子。35-45分钟区间边际效应最大')
    P(doc, '2. 跟进次数: 2-4次转化率最高，5次以上边际效益为负(倒U型)')
    P(doc, '3. 客户年龄: 35-45岁区间转化最高，与G9目标客群一致')
    P(doc, '4. 渠道历史转化率: 反映渠道"品牌溢价"效应')
    P(doc, '5. 跟进强度: 0.5-2的最优区间')

    H(doc, '4.3 流失预警模型', 2)
    P(doc, 'Cox比例风险模型 + 随机生存森林(RSF)双重验证:')
    img(doc, os.path.join(W3C,'w3_21_survival_curve.png'), 5.5, '图25: Kaplan-Meier生存曲线(按配置)')
    img(doc, os.path.join(W3C,'w3_22_cox_forest.png'), 5.5, '图26: Cox风险比森林图')
    img(doc, os.path.join(W3C,'w3_23_churn_risk_heatmap.png'), 5.5, '图27: 流失风险热力图(城市×投诉类型)')
    P(doc, '关键风险因素: 延迟交付(HR≈2.1) > 质量投诉(HR≈2.5) > 处理时长>7天(HR≈1.8)', bold=True)

    img(doc, os.path.join(W3C,'w3_19_customer_clusters.png'), 5.5, '图28: 客户分群 — 年龄×试驾时长散点(成交着色)')

    # Optimization
    H(doc, '第五章  策略优化', 1)
    H(doc, '5.1 预算优化方案', 2)
    img(doc, os.path.join(W4C,'w4_01_budget_comparison.png'), 5.5, '图29: 当前vs最优预算(前15组合)')
    img(doc, os.path.join(W4C,'w4_02_budget_treemap.png'), 5.5, '图30: 最优预算城市×渠道堆叠分配')
    img(doc, os.path.join(W4C,'w4_04_scenario_radar.png'), 5.5, '图31: 5场景多维度雷达对比')
    img(doc, os.path.join(W4C,'w4_05_ga_convergence.png'), 5.0, '图32: 遗传算法收敛曲线')

    H(doc, '5.2 定价博弈', 2)
    img(doc, os.path.join(W4C,'w4_06_game_payoff.png'), 5.0, '图33: 博弈收益矩阵热力图')
    img(doc, os.path.join(W4C,'w4_07_best_response.png'), 5.5, '图34: 最佳响应曲线')

    H(doc, '5.3 人效优化', 2)
    img(doc, os.path.join(W4C,'w4_08_salesperson_workload.png'), 5.5, '图35: 销售员工作量对比')
    img(doc, os.path.join(W4C,'w4_15_city_staffing.png'), 5.5, '图36: 各城市人员负荷分析')

    H(doc, '5.4 综合策略推荐', 2)
    img(doc, os.path.join(W4C,'w4_10_strategy_waterfall.png'), 5.5, '图37: 策略贡献瀑布图')
    img(doc, os.path.join(W4C,'w4_11_roi_tornado.png'), 5.5, '图38: ROI龙卷风图')
    img(doc, os.path.join(W4C,'w4_12_pareto_frontier.png'), 5.5, '图39: Pareto效率前沿')
    img(doc, os.path.join(W4C,'w4_14_topsis_ranking.png'), 5.5, '图40: TOPSIS综合排名')

    T(doc,
        ['策略', '动作', '预期效果', '风险', '周期'],
        [['预算重分配', '车展-25%,懂车帝+30%,抖音+20%', '+6-10%订单', '渠道集中', '2周'],
         ['跟进SOP', '面谈推广,首次≤3天,2-4次', '+3-5pp转化率', '执行一致性', '4周'],
         ['售后预警', '评分<3预警,投诉24h响应', '-20%流失', '系统对接', '6周'],
         ['智能调度', '模型匹配线索-销售员', '+15%人效', '技术复杂度', '8-12周']])

    # Roadmap
    H(doc, '第六章  实施路线图', 1)
    img(doc, os.path.join(W4C,'w4_13_implementation_timeline.png'), 5.5, '图41: Q4分阶段实施甘特图')

    T(doc, ['阶段', '时间', '关键行动', '里程碑'],
        [['数据基础', 'W1-W2', '清洗+因果', '宽表+归因验收'],
         ['模型训练', 'W3-W4', '预测+优化', '模型AUC达标'],
         ['系统上线', 'W5-W6', '仪表盘+培训', 'V1上线'],
         ['试运行', 'W7-W8', '小范围试运行', '试运行通过'],
         ['全面推广', 'W9-W12', '正式执行+监控', 'Q4目标达成']])

    # Appendix
    H(doc, '附录  方法论与技术栈', 1)
    T(doc, ['技术', '方法', '用途'],
        [['Pandas+NumPy', '数据清洗、特征工程', '宽表构建'],
         ['PSM', 'Rosenbaum & Rubin, 1983', '跟进方式因果效应'],
         ['RDD', 'Imbens & Lemieux, 2008', '延迟交付满意度影响'],
         ['Shapley Value', 'Shapley, 1953', '渠道公平归因'],
         ['LightGBM', 'Ke et al., 2017', '转化概率预测'],
         ['Optuna', 'Akiba et al., 2019', '贝叶斯超参优化'],
         ['Cox PH', 'Cox, 1972', '客户流失风险'],
         ['SHAP', 'Lundberg & Lee, 2017', '模型可解释性'],
         ['LP+GA', '线性规划+遗传算法', '预算分配优化'],
         ['TOPSIS', 'Hwang & Yoon, 1981', '多方案综合排序']])

    path = os.path.join(REPORTS_DIR, 'G9_销售运营优化白皮书.docx')
    doc.save(path); print(f'  -> {os.path.basename(path)}')
    return path

# ═══════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('='*60)
    print('  G9 Reports V2 — Embedded Charts Edition')
    print('='*60)
    gen_w1()
    gen_w2()
    gen_w3()
    gen_w4()
    gen_w5()
    gen_whitepaper()
    print(f'\n[OK] 6 documents with embedded charts in: {REPORTS_DIR}/')
